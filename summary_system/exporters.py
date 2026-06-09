from __future__ import annotations

import json
from pathlib import Path
from typing import Any

import fitz
from docx import Document
from docx.shared import Pt

from .models import SummaryResult


def result_to_markdown(result: SummaryResult) -> str:
    lines = [
        f"# {result.title}",
        "",
        f"- 场景类型：{result.scene.scene_type}",
        f"- 置信度：{result.scene.confidence}",
        f"- 推荐模板：{result.scene.recommended_template}",
        f"- 生成时间：{result.created_at}",
        "",
        "## 识别依据",
    ]
    lines.extend(f"- {reason}" for reason in result.scene.reasons)
    lines.extend(["", "## 总结内容"])
    lines.extend(_value_to_markdown(result.content, 3))
    lines.extend(["", "## 质量校验"])
    if result.quality_issues:
        lines.extend(f"- [{issue.level}] {issue.message}" for issue in result.quality_issues)
    else:
        lines.append("- 未发现明显质量问题。")
    return "\n".join(lines) + "\n"


def _value_to_markdown(value: Any, heading_level: int) -> list[str]:
    lines: list[str] = []
    if isinstance(value, dict):
        for key, item in value.items():
            lines.append(f"{'#' * heading_level} {key}")
            lines.extend(_value_to_markdown(item, min(heading_level + 1, 6)))
            lines.append("")
    elif isinstance(value, list):
        for item in value:
            if isinstance(item, dict):
                lines.append("- " + "；".join(f"{k}：{v}" for k, v in item.items()))
            else:
                lines.append(f"- {item}")
    else:
        lines.append(str(value))
    return lines


def save_markdown(result: SummaryResult, path: Path) -> Path:
    path.write_text(result_to_markdown(result), encoding="utf-8")
    return path


def save_json(result: SummaryResult, path: Path) -> Path:
    from .agents import AgentOrchestrator

    path.write_text(json.dumps(AgentOrchestrator.to_dict(result), ensure_ascii=False, indent=2), encoding="utf-8")
    return path


def save_docx(result: SummaryResult, path: Path) -> Path:
    doc = Document()
    doc.styles["Normal"].font.name = "Microsoft YaHei"
    doc.styles["Normal"].font.size = Pt(10.5)
    doc.add_heading(result.title, level=0)
    doc.add_paragraph(f"场景类型：{result.scene.scene_type}    置信度：{result.scene.confidence}")
    doc.add_paragraph(f"推荐模板：{result.scene.recommended_template}    生成时间：{result.created_at}")
    doc.add_heading("识别依据", level=1)
    for reason in result.scene.reasons:
        doc.add_paragraph(reason, style="List Bullet")
    doc.add_heading("总结内容", level=1)
    _append_docx_value(doc, result.content, 2)
    doc.add_heading("质量校验", level=1)
    if result.quality_issues:
        for issue in result.quality_issues:
            doc.add_paragraph(f"[{issue.level}] {issue.message}", style="List Bullet")
    else:
        doc.add_paragraph("未发现明显质量问题。")
    doc.save(path)
    return path


def _append_docx_value(doc: Document, value: Any, level: int) -> None:
    if isinstance(value, dict):
        for key, item in value.items():
            doc.add_heading(str(key), level=min(level, 4))
            _append_docx_value(doc, item, min(level + 1, 4))
    elif isinstance(value, list):
        if value and all(isinstance(item, dict) for item in value):
            keys = list(value[0].keys())
            table = doc.add_table(rows=1, cols=len(keys))
            table.style = "Table Grid"
            for idx, key in enumerate(keys):
                table.rows[0].cells[idx].text = str(key)
            for item in value:
                cells = table.add_row().cells
                for idx, key in enumerate(keys):
                    cells[idx].text = str(item.get(key, ""))
        else:
            for item in value:
                doc.add_paragraph(str(item), style="List Bullet")
    else:
        doc.add_paragraph(str(value))


def save_pdf(result: SummaryResult, path: Path) -> Path:
    markdown = result_to_markdown(result)
    pdf = fitz.open()
    fontfile = _find_chinese_font()
    page = pdf.new_page(width=595, height=842)
    rect = fitz.Rect(50, 50, 545, 792)
    fontname = "chinafont" if fontfile else "helv"
    if fontfile:
        page.insert_font(fontname=fontname, fontfile=str(fontfile))
    y = rect.y0
    for raw_line in markdown.splitlines():
        line = raw_line.strip()
        if not line:
            y += 10
            continue
        size = 15 if line.startswith("# ") else 12 if line.startswith("##") else 10
        text = line.lstrip("#").strip()
        for part in _wrap_text(text, 38 if size >= 12 else 48):
            if y > rect.y1 - 24:
                page = pdf.new_page(width=595, height=842)
                if fontfile:
                    page.insert_font(fontname=fontname, fontfile=str(fontfile))
                y = rect.y0
            page.insert_text((rect.x0, y), part, fontsize=size, fontname=fontname)
            y += size + 7
    pdf.save(path)
    pdf.close()
    return path


def _wrap_text(text: str, width: int) -> list[str]:
    if len(text) <= width:
        return [text]
    return [text[idx : idx + width] for idx in range(0, len(text), width)]


def _find_chinese_font() -> Path | None:
    candidates = [
        Path(r"C:\Windows\Fonts\msyh.ttc"),
        Path(r"C:\Windows\Fonts\simsun.ttc"),
        Path(r"C:\Windows\Fonts\simhei.ttf"),
    ]
    return next((path for path in candidates if path.exists()), None)

